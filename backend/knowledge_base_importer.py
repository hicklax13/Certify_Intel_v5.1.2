"""
Certify Intel - Knowledge Base Importer (v5.0.8)

Imports competitor data from client-provided files in the knowledge base folder.
Supports CSV, Excel, PDF, Word, and Markdown files.

All imported data is labeled as "Source: Certify Health" with source_type="client_provided".
"""

import os
import csv
import json
import re
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, asdict, field
from pathlib import Path

# Excel parsing
try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    print("[KB Importer] openpyxl not available - Excel parsing disabled")

# Word document parsing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[KB Importer] python-docx not available - Word parsing disabled")

# PDF parsing (optional - can use Gemini instead)
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("[KB Importer] PyMuPDF not available - will use Gemini for PDFs")


# ==============================================================================
# COMPETITOR NAME NORMALIZATION
# ==============================================================================

# Master list of 74 competitors with aliases for matching
COMPETITOR_ALIASES = {
    "1st provider's choice": ["1st provider's choice", "1st providers choice", "first provider's choice"],
    "access": ["access"],
    "advancedmd": ["advancedmd", "advanced md", "advanced-md"],
    "advantech": ["advantech", "advan tech"],
    "aila tech": ["aila tech", "ailatech", "aila"],
    "athenahealth": ["athenahealth", "athena health", "athena"],
    "axiamed": ["axiamed", "axia med"],
    "bestdocapp": ["bestdocapp", "best doc app"],
    "bridgeinteract": ["bridgeinteract", "bridge interact"],
    "carecloud": ["carecloud", "carecloud - breeze", "carecloud breeze", "care cloud"],
    "centralreach": ["centralreach", "central reach"],
    "change healthcare": ["change healthcare", "changehealthcare"],
    "cipherhealth": ["cipherhealth", "cipher health"],
    "clearwave": ["clearwave", "clear wave"],
    "cliexa": ["cliexa"],
    "cliniko": ["cliniko"],
    "compugroup medical": ["compugroup medical", "compugroup medical (cgm)", "cgm", "compugroup"],
    "creo": ["creo"],
    "curago health": ["curago health", "curago"],
    "curemd": ["curemd", "cure md"],
    "doctible": ["doctible"],
    "drchrono": ["drchrono", "dr chrono"],
    "eloTouch": ["elotouch", "elo touch"],
    "epion health": ["epion health", "epion", "kyruus/epion"],
    "ezreferral": ["ezreferral", "ez referral"],
    "followmyhealth": ["followmyhealth", "follow my health"],
    "formdr": ["formdr", "form dr"],
    "formstack": ["formstack", "form stack"],
    "frank mayer": ["frank mayer"],
    "get well": ["get well", "getwell"],
    "greyfinch": ["greyfinch", "grey finch"],
    "health asyst": ["health asyst", "health asyst - checkin asyst", "checkin asyst", "healthasyst"],
    "health ipass": ["health ipass", "healthipass", "health-ipass"],
    "healthmark group": ["healthmark group", "healthmark"],
    "imprivata": ["imprivata"],
    "insynchus": ["insynchus", "in synchus"],
    "intakeq": ["intakeq", "intake q"],
    "intelichart": ["intelichart", "inteli chart"],
    "interlace health": ["interlace health", "interlace"],
    "intiveo": ["intiveo"],
    "intouch with health": ["intouch with health", "intouch health", "intouch"],
    "intrado": ["intrado"],
    "jellyfish health": ["jellyfish health", "jellyfish health acquired by millennia", "millennia", "jellyfish"],
    "kareo": ["kareo", "kareo/tebra", "tebra"],
    "kyruus": ["kyruus"],
    "lumahealth": ["lumahealth", "luma health", "luma"],
    "mconsent": ["mconsent", "m consent"],
    "medfusion": ["medfusion", "med fusion"],
    "mend": ["mend"],
    "nexhealth": ["nexhealth", "nex health"],
    "nextech": ["nextech", "nex tech"],
    "notable health": ["notable health", "notable"],
    "orbitia": ["orbitia"],
    "patient communicator": ["patient communicator", "patientcommunicator"],
    "patientpop": ["patientpop", "patient pop"],
    "patienttrak": ["patienttrak", "patient trak"],
    "patientworks": ["patientworks", "patient works"],
    "phreesia": ["phreesia"],
    "pomelo health": ["pomelo health", "pomelo"],
    "quadramed": ["quadramed", "quadra med"],
    "relatient": ["relatient"],
    "right patient": ["right patient", "rightpatient"],
    "simplepractice": ["simplepractice", "simple practice"],
    "solutionreach": ["solutionreach", "solution reach"],
    "televox": ["televox", "tele vox"],
    "tonic": ["tonic", "tonic health", "tonicforhealth"],
    "touchhealth": ["touchhealth", "touch health"],
    "triassic solutions": ["triassic solutions", "triassic"],
    "updox": ["updox"],
    "vecnahealth": ["vecnahealth", "vecna health", "vecna"],
    "weave": ["weave"],
    "well health": ["well health", "wellhealth"],
    "yapiaapp": ["yapiaapp", "yapia app", "yapia"],
}

# Reverse lookup: alias -> canonical name
ALIAS_TO_CANONICAL = {}
for canonical, aliases in COMPETITOR_ALIASES.items():
    for alias in aliases:
        ALIAS_TO_CANONICAL[alias.lower().strip()] = canonical


def normalize_competitor_name(name: str) -> str:
    """
    Normalize competitor name for matching.

    Args:
        name: Raw competitor name from source file

    Returns:
        Canonical competitor name or original if no match
    """
    if not name:
        return ""

    normalized = name.lower().strip()

    # Direct alias lookup
    if normalized in ALIAS_TO_CANONICAL:
        return ALIAS_TO_CANONICAL[normalized]

    # Fuzzy matching: remove special chars and try again
    cleaned = re.sub(r'[^\w\s]', '', normalized)
    if cleaned in ALIAS_TO_CANONICAL:
        return ALIAS_TO_CANONICAL[cleaned]

    # Try matching first part of name
    first_word = normalized.split()[0] if normalized.split() else normalized
    if first_word in ALIAS_TO_CANONICAL:
        return ALIAS_TO_CANONICAL[first_word]

    # Return original (capitalized) if no match found
    return name.strip()


# ==============================================================================
# DATA STRUCTURES
# ==============================================================================

@dataclass
class FileInfo:
    """Information about a file in the knowledge base."""
    path: str
    filename: str
    extension: str
    size_bytes: int
    file_type: str  # csv, xlsx, pdf, docx, md, txt


@dataclass
class CompetitorData:
    """Extracted competitor data from knowledge base files."""
    name: str
    canonical_name: str  # Normalized name
    source_file: str
    source_type: str = "client_provided"

    # Basic info
    website: Optional[str] = None
    status: str = "Active"
    threat_level: str = "Medium"

    # Pricing
    pricing_model: Optional[str] = None
    base_price: Optional[str] = None
    price_unit: Optional[str] = None

    # Product
    product_categories: Optional[str] = None
    key_features: Optional[str] = None
    integration_partners: Optional[str] = None
    certifications: Optional[str] = None

    # Market
    target_segments: Optional[str] = None
    customer_size_focus: Optional[str] = None
    geographic_focus: Optional[str] = None
    customer_count: Optional[str] = None
    customer_acquisition_rate: Optional[str] = None
    key_customers: Optional[str] = None

    # Company
    employee_count: Optional[str] = None
    employee_growth_rate: Optional[str] = None
    year_founded: Optional[str] = None
    headquarters: Optional[str] = None
    funding_total: Optional[str] = None
    latest_round: Optional[str] = None
    pe_vc_backers: Optional[str] = None

    # Digital
    website_traffic: Optional[str] = None
    social_following: Optional[str] = None
    g2_rating: Optional[str] = None
    capterra_rating: Optional[str] = None

    # Features (boolean flags)
    has_patient_intake: bool = False
    has_insurance_verification: bool = False
    has_payments: bool = False
    has_scheduling: bool = False
    has_patient_portal: bool = False

    # Notes
    notes: Optional[str] = None
    differentiators: Optional[str] = None
    recent_launches: Optional[str] = None

    # Additional fields extracted
    extra_fields: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ImportResult:
    """Result of importing knowledge base files."""
    total_files_scanned: int
    files_parsed: int
    competitors_found: int
    unique_competitors: int
    competitors: List[CompetitorData]
    errors: List[str]
    warnings: List[str]


# ==============================================================================
# CSV COLUMN MAPPING
# ==============================================================================

# Map CSV column headers to CompetitorData fields
CSV_COLUMN_MAPPING = {
    # Basic info
    "competitor name": "name",
    "name": "name",
    "company": "name",
    "website": "website",
    "url": "website",
    "status": "status",
    "threat level": "threat_level",
    "threat": "threat_level",

    # Pricing
    "pricing model": "pricing_model",
    "pricing": "pricing_model",
    "base price": "base_price",
    "price": "base_price",
    "price unit": "price_unit",

    # Product
    "product categories": "product_categories",
    "products": "product_categories",
    "categories": "product_categories",
    "key features": "key_features",
    "features": "key_features",
    "integration partners": "integration_partners",
    "integrations": "integration_partners",
    "certifications": "certifications",

    # Market
    "target segments": "target_segments",
    "segments": "target_segments",
    "customer size focus": "customer_size_focus",
    "size focus": "customer_size_focus",
    "geographic focus": "geographic_focus",
    "geo": "geographic_focus",
    "customer count": "customer_count",
    "customers": "customer_count",
    "customer acquisition rate": "customer_acquisition_rate",
    "key customers": "key_customers",

    # Company
    "employee count": "employee_count",
    "employees": "employee_count",
    "employee growth rate": "employee_growth_rate",
    "year founded": "year_founded",
    "founded": "year_founded",
    "hq location": "headquarters",
    "headquarters": "headquarters",
    "hq": "headquarters",
    "funding total": "funding_total",
    "funding": "funding_total",
    "latest round": "latest_round",
    "pe/vc backers": "pe_vc_backers",
    "investors": "pe_vc_backers",

    # Digital
    "website traffic (monthly)": "website_traffic",
    "website traffic": "website_traffic",
    "traffic": "website_traffic",
    "social media following": "social_following",
    "social following": "social_following",
    "g2 rating": "g2_rating",
    "g2": "g2_rating",
    "capterra rating": "capterra_rating",
    "capterra": "capterra_rating",

    # Features (booleans)
    "feature: patient intake": "has_patient_intake",
    "patient intake": "has_patient_intake",
    "feature: insurance verification": "has_insurance_verification",
    "insurance verification": "has_insurance_verification",
    "feature: payments": "has_payments",
    "payments": "has_payments",
    "feature: scheduling": "has_scheduling",
    "scheduling": "has_scheduling",
    "feature: patient portal": "has_patient_portal",
    "patient portal": "has_patient_portal",

    # Notes
    "notes": "notes",
    "differentiators": "differentiators",
    "recent product launches": "recent_launches",
    "recent launches": "recent_launches",
    "tagline": "extra_fields.tagline",
}


# ==============================================================================
# KNOWLEDGE BASE IMPORTER
# ==============================================================================

class KnowledgeBaseImporter:
    """
    Imports competitor data from knowledge base files.

    Supports:
    - CSV files (direct column mapping)
    - Excel files (multi-sheet parsing)
    - PDF files (Gemini multimodal or PyMuPDF)
    - Word documents (python-docx)
    - Markdown files (text parsing)
    """

    def __init__(self, folder_path: str = None):
        """
        Initialize the importer.

        Args:
            folder_path: Path to knowledge base folder.
                         Defaults to project's client_docs_knowledge_base folder.
        """
        if folder_path is None:
            # Default to project's knowledge base folder
            project_root = Path(__file__).parent.parent
            folder_path = project_root / "client_docs_knowledge_base"

        self.folder_path = Path(folder_path)
        self.errors: List[str] = []
        self.warnings: List[str] = []
        self.gemini_provider = None

        # Try to load Gemini provider for PDF analysis
        try:
            from gemini_provider import GeminiProvider
            self.gemini_provider = GeminiProvider()
        except ImportError:
            print("[KB Importer] Gemini provider not available for PDF analysis")

    def scan_folder(self) -> List[FileInfo]:
        """
        Scan the knowledge base folder for supported files.

        Returns:
            List of FileInfo objects for each supported file
        """
        files = []

        if not self.folder_path.exists():
            self.errors.append(f"Folder not found: {self.folder_path}")
            return files

        # Walk through folder and subfolders
        for root, dirs, filenames in os.walk(self.folder_path):
            # Skip copy folders to avoid duplicates
            if "copy" in root.lower() and "- copy" in root.lower():
                continue

            for filename in filenames:
                filepath = Path(root) / filename
                ext = filepath.suffix.lower()

                # Determine file type
                file_type = None
                if ext == ".csv":
                    file_type = "csv"
                elif ext in [".xlsx", ".xls"]:
                    file_type = "xlsx"
                elif ext == ".pdf":
                    file_type = "pdf"
                elif ext in [".docx", ".doc"]:
                    file_type = "docx"
                elif ext == ".md":
                    file_type = "md"
                elif ext == ".txt":
                    file_type = "txt"

                if file_type:
                    files.append(FileInfo(
                        path=str(filepath),
                        filename=filename,
                        extension=ext,
                        size_bytes=filepath.stat().st_size,
                        file_type=file_type
                    ))

        return files

    def parse_csv(self, filepath: str) -> List[CompetitorData]:
        """
        Parse a CSV file for competitor data.

        Args:
            filepath: Path to CSV file

        Returns:
            List of CompetitorData objects
        """
        competitors = []

        try:
            with open(filepath, 'r', encoding='utf-8-sig') as f:
                reader = csv.DictReader(f)

                for row in reader:
                    # Skip empty rows
                    if not any(row.values()):
                        continue

                    # Find the name column
                    name = None
                    for col_name, value in row.items():
                        col_lower = col_name.lower().strip()
                        if col_lower in ["competitor name", "name", "company"]:
                            name = value.strip() if value else None
                            break

                    if not name:
                        continue

                    # Create competitor data
                    competitor = CompetitorData(
                        name=name,
                        canonical_name=normalize_competitor_name(name),
                        source_file=filepath
                    )

                    # Map columns to fields
                    for col_name, value in row.items():
                        if not value or not value.strip():
                            continue

                        col_lower = col_name.lower().strip()

                        if col_lower in CSV_COLUMN_MAPPING:
                            field_name = CSV_COLUMN_MAPPING[col_lower]

                            # Handle boolean fields
                            if field_name.startswith("has_"):
                                setattr(competitor, field_name,
                                       value.lower().strip() in ["yes", "true", "1", "y"])
                            # Handle extra fields
                            elif field_name.startswith("extra_fields."):
                                key = field_name.split(".", 1)[1]
                                competitor.extra_fields[key] = value.strip()
                            else:
                                setattr(competitor, field_name, value.strip())

                    competitors.append(competitor)

        except Exception as e:
            self.errors.append(f"Error parsing CSV {filepath}: {str(e)}")

        return competitors

    def parse_excel(self, filepath: str) -> List[CompetitorData]:
        """
        Parse an Excel file for competitor data.

        Args:
            filepath: Path to Excel file

        Returns:
            List of CompetitorData objects
        """
        if not OPENPYXL_AVAILABLE:
            self.warnings.append(f"Skipping Excel file (openpyxl not installed): {filepath}")
            return []

        competitors = []

        try:
            wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]

                # Get header row
                headers = []
                for cell in sheet[1]:
                    headers.append(str(cell.value).lower().strip() if cell.value else "")

                # Find name column index
                name_col_idx = None
                for i, header in enumerate(headers):
                    if header in ["competitor name", "name", "company", "competitor"]:
                        name_col_idx = i
                        break

                if name_col_idx is None:
                    continue  # Skip sheets without name column

                # Parse data rows
                for row_idx, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
                    if not row or not row[name_col_idx]:
                        continue

                    name = str(row[name_col_idx]).strip()
                    if not name:
                        continue

                    competitor = CompetitorData(
                        name=name,
                        canonical_name=normalize_competitor_name(name),
                        source_file=f"{filepath}:{sheet_name}"
                    )

                    # Map columns to fields
                    for col_idx, value in enumerate(row):
                        if col_idx >= len(headers) or not value:
                            continue

                        col_name = headers[col_idx]
                        value_str = str(value).strip()

                        if col_name in CSV_COLUMN_MAPPING:
                            field_name = CSV_COLUMN_MAPPING[col_name]

                            if field_name.startswith("has_"):
                                setattr(competitor, field_name,
                                       value_str.lower() in ["yes", "true", "1", "y"])
                            elif field_name.startswith("extra_fields."):
                                key = field_name.split(".", 1)[1]
                                competitor.extra_fields[key] = value_str
                            else:
                                setattr(competitor, field_name, value_str)

                    competitors.append(competitor)

            wb.close()

        except Exception as e:
            self.errors.append(f"Error parsing Excel {filepath}: {str(e)}")

        return competitors

    def parse_pdf(self, filepath: str) -> List[CompetitorData]:
        """
        Parse a PDF file for competitor data using Gemini multimodal.

        Args:
            filepath: Path to PDF file

        Returns:
            List of CompetitorData objects
        """
        competitors = []

        # Try Gemini multimodal first (preferred)
        if self.gemini_provider:
            try:
                result = self._parse_pdf_with_gemini(filepath)
                if result:
                    competitors.extend(result)
                    return competitors
            except Exception as e:
                self.warnings.append(f"Gemini PDF parsing failed for {filepath}: {str(e)}")

        # Fallback to PyMuPDF text extraction
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(filepath)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()

                # Extract competitor names from text
                competitors.extend(self._extract_competitors_from_text(text, filepath))
            except Exception as e:
                self.errors.append(f"Error parsing PDF {filepath}: {str(e)}")
        else:
            self.warnings.append(f"Skipping PDF file (no parser available): {filepath}")

        return competitors

    def _parse_pdf_with_gemini(self, filepath: str) -> List[CompetitorData]:
        """Parse PDF using Gemini multimodal API."""
        if not self.gemini_provider:
            return []

        prompt = """
        Analyze this competitive intelligence document and extract all competitor information.
        For each competitor mentioned, extract the following if available:
        - Company name
        - Website
        - Pricing model and price
        - Product categories/features
        - Target market/segments
        - Employee count
        - Funding/valuation
        - Strengths and weaknesses
        - Key differentiators

        Return a JSON array of competitors with these fields.
        If no competitors are found, return an empty array [].
        """

        try:
            # Use Gemini's PDF analysis capability
            result = self.gemini_provider.analyze_pdf_sync(filepath, prompt)

            if result and isinstance(result, list):
                competitors = []
                for item in result:
                    if isinstance(item, dict) and item.get("name"):
                        competitor = CompetitorData(
                            name=item.get("name", ""),
                            canonical_name=normalize_competitor_name(item.get("name", "")),
                            source_file=filepath,
                            website=item.get("website"),
                            pricing_model=item.get("pricing_model"),
                            base_price=item.get("price"),
                            product_categories=item.get("products"),
                            target_segments=item.get("target_market"),
                            employee_count=item.get("employees"),
                            funding_total=item.get("funding"),
                            differentiators=item.get("differentiators"),
                        )
                        competitors.append(competitor)
                return competitors
        except Exception as e:
            self.warnings.append(f"Gemini PDF analysis failed: {str(e)}")

        return []

    def _extract_competitors_from_text(self, text: str, source_file: str) -> List[CompetitorData]:
        """Extract competitor names from text by matching against known names."""
        competitors = []
        found_names = set()

        # Look for known competitor names in text
        text_lower = text.lower()

        for canonical, aliases in COMPETITOR_ALIASES.items():
            for alias in aliases:
                if alias in text_lower and canonical not in found_names:
                    found_names.add(canonical)
                    competitors.append(CompetitorData(
                        name=canonical.title(),
                        canonical_name=canonical,
                        source_file=source_file
                    ))
                    break

        return competitors

    def parse_docx(self, filepath: str) -> List[CompetitorData]:
        """
        Parse a Word document for competitor data.

        Args:
            filepath: Path to Word document

        Returns:
            List of CompetitorData objects
        """
        if not DOCX_AVAILABLE:
            self.warnings.append(f"Skipping Word file (python-docx not installed): {filepath}")
            return []

        competitors = []

        try:
            doc = DocxDocument(filepath)

            # Extract all text
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)

            text = "\n".join(full_text)

            # Extract competitors from text
            competitors.extend(self._extract_competitors_from_text(text, filepath))

            # Special handling for battlecard documents
            if "battlecard" in filepath.lower():
                competitors.extend(self._parse_battlecard_docx(doc, filepath))

        except Exception as e:
            self.errors.append(f"Error parsing Word doc {filepath}: {str(e)}")

        return competitors

    def _parse_battlecard_docx(self, doc, filepath: str) -> List[CompetitorData]:
        """Extract structured battlecard data from Word document."""
        competitors = []
        current_competitor = None
        current_section = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if this is a competitor heading
            normalized = normalize_competitor_name(text)
            if normalized in COMPETITOR_ALIASES:
                if current_competitor:
                    competitors.append(current_competitor)
                current_competitor = CompetitorData(
                    name=text,
                    canonical_name=normalized,
                    source_file=filepath
                )
                current_section = None
                continue

            # Check for section headers
            text_lower = text.lower()
            if "strength" in text_lower:
                current_section = "strengths"
            elif "weakness" in text_lower:
                current_section = "weaknesses"
            elif "differentiator" in text_lower:
                current_section = "differentiators"
            elif current_competitor and current_section:
                # Add content to current section
                if current_section == "differentiators":
                    existing = current_competitor.differentiators or ""
                    current_competitor.differentiators = (existing + "\n" + text).strip()

        if current_competitor:
            competitors.append(current_competitor)

        return competitors

    def parse_markdown(self, filepath: str) -> List[CompetitorData]:
        """
        Parse a Markdown file for competitor data.

        Args:
            filepath: Path to Markdown file

        Returns:
            List of CompetitorData objects
        """
        competitors = []

        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()

            # Extract competitors from text
            competitors.extend(self._extract_competitors_from_text(text, filepath))

        except Exception as e:
            self.errors.append(f"Error parsing Markdown {filepath}: {str(e)}")

        return competitors

    def extract_all(self) -> ImportResult:
        """
        Extract competitor data from all files in the knowledge base folder.

        Returns:
            ImportResult with all extracted competitors and statistics
        """
        all_competitors: List[CompetitorData] = []
        files = self.scan_folder()
        files_parsed = 0

        for file_info in files:
            try:
                if file_info.file_type == "csv":
                    competitors = self.parse_csv(file_info.path)
                elif file_info.file_type == "xlsx":
                    competitors = self.parse_excel(file_info.path)
                elif file_info.file_type == "pdf":
                    competitors = self.parse_pdf(file_info.path)
                elif file_info.file_type == "docx":
                    competitors = self.parse_docx(file_info.path)
                elif file_info.file_type in ["md", "txt"]:
                    competitors = self.parse_markdown(file_info.path)
                else:
                    continue

                all_competitors.extend(competitors)
                files_parsed += 1

            except Exception as e:
                self.errors.append(f"Error processing {file_info.path}: {str(e)}")

        # Deduplicate by canonical name
        unique_competitors = self._deduplicate_competitors(all_competitors)

        return ImportResult(
            total_files_scanned=len(files),
            files_parsed=files_parsed,
            competitors_found=len(all_competitors),
            unique_competitors=len(unique_competitors),
            competitors=unique_competitors,
            errors=self.errors,
            warnings=self.warnings
        )

    def _deduplicate_competitors(self, competitors: List[CompetitorData]) -> List[CompetitorData]:
        """
        Deduplicate competitors by canonical name, merging data from multiple sources.

        Args:
            competitors: List of all extracted competitors (may have duplicates)

        Returns:
            List of unique competitors with merged data
        """
        merged: Dict[str, CompetitorData] = {}

        for comp in competitors:
            canonical = comp.canonical_name.lower()

            if canonical not in merged:
                merged[canonical] = comp
            else:
                # Merge: fill gaps in existing with new data
                existing = merged[canonical]

                for field_name in [
                    "website", "pricing_model", "base_price", "price_unit",
                    "product_categories", "key_features", "integration_partners",
                    "certifications", "target_segments", "customer_size_focus",
                    "geographic_focus", "customer_count", "customer_acquisition_rate",
                    "key_customers", "employee_count", "employee_growth_rate",
                    "year_founded", "headquarters", "funding_total", "latest_round",
                    "pe_vc_backers", "website_traffic", "social_following",
                    "g2_rating", "capterra_rating", "notes", "differentiators",
                    "recent_launches"
                ]:
                    existing_val = getattr(existing, field_name, None)
                    new_val = getattr(comp, field_name, None)

                    # Fill gaps only (don't overwrite existing data)
                    if not existing_val and new_val:
                        setattr(existing, field_name, new_val)

                # Merge boolean flags with OR
                for bool_field in ["has_patient_intake", "has_insurance_verification",
                                   "has_payments", "has_scheduling", "has_patient_portal"]:
                    if getattr(comp, bool_field, False):
                        setattr(existing, bool_field, True)

                # Track multiple source files
                if comp.source_file not in existing.source_file:
                    existing.source_file = f"{existing.source_file}; {comp.source_file}"

        return list(merged.values())


# ==============================================================================
# UTILITY FUNCTIONS
# ==============================================================================

def get_all_competitor_names() -> List[str]:
    """Get list of all 74 canonical competitor names."""
    return list(COMPETITOR_ALIASES.keys())


def is_known_competitor(name: str) -> bool:
    """Check if a name matches a known competitor."""
    normalized = normalize_competitor_name(name)
    return normalized.lower() in COMPETITOR_ALIASES


# ==============================================================================
# CLI TESTING
# ==============================================================================

# ==============================================================================
# PREINSTALL FUNCTION - Auto-import on first startup
# ==============================================================================

def preinstall_knowledge_base(db_session=None):
    """
    Auto-import knowledge base on first startup.

    This function is called during server startup to ensure client-provided
    competitor data is preinstalled in the database. It only runs once -
    subsequent startups skip the import.

    Features:
    - Runs only on first startup (tracks via SystemSetting)
    - Labels all data as "Certify Health (Preinstalled)"
    - Uses fill-gaps-only mode (never overwrites existing data)
    - Creates DataSource records for tracking
    - Logs the import activity

    Args:
        db_session: Optional SQLAlchemy session. If None, creates its own.

    Returns:
        dict with import results or None if skipped
    """
    from database import SessionLocal, Competitor, DataSource, SystemSetting

    # Create session if not provided
    if db_session is None:
        db = SessionLocal()
        close_db = True
    else:
        db = db_session
        close_db = False

    try:
        # Check if preinstall has already been done
        setting = db.query(SystemSetting).filter(
            SystemSetting.key == "knowledge_base_preinstalled"
        ).first()

        if setting and setting.value == "true":
            print("[KB Preinstall] Already preinstalled, skipping...")
            return None

        print("[KB Preinstall] Starting preinstall of client-provided knowledge base...")

        # Run the import
        importer = KnowledgeBaseImporter()
        result = importer.extract_all()

        if result.competitors_found == 0:
            print("[KB Preinstall] No competitors found in knowledge base folder")
            # Still mark as done to avoid re-checking empty folder
            _mark_preinstall_complete(db)
            if close_db:
                db.close()
            return {"skipped": True, "reason": "no_competitors_found"}

        # Import competitors
        imported_count = 0
        updated_count = 0

        for comp_data in result.competitors:
            try:
                # Check if competitor already exists
                existing = db.query(Competitor).filter(
                    Competitor.name.ilike(comp_data.canonical_name)
                ).first()

                if existing:
                    # Update existing (fill gaps only)
                    fields_updated = _preinstall_update_competitor(
                        db, existing, comp_data
                    )
                    if fields_updated > 0:
                        updated_count += 1
                else:
                    # Create new competitor
                    _preinstall_create_competitor(db, comp_data)
                    imported_count += 1

            except Exception as e:
                print(f"[KB Preinstall] Error importing {comp_data.canonical_name}: {e}")

        # Commit all changes
        db.commit()

        # Mark preinstall as complete
        _mark_preinstall_complete(db)

        print(f"[KB Preinstall] Complete! {imported_count} new, {updated_count} updated")
        print(f"[KB Preinstall] Data labeled as 'Certify Health (Preinstalled)'")

        return {
            "success": True,
            "competitors_imported": imported_count,
            "competitors_updated": updated_count,
            "total_found": result.unique_competitors
        }

    except Exception as e:
        print(f"[KB Preinstall] Error: {e}")
        db.rollback()
        return {"success": False, "error": str(e)}

    finally:
        if close_db:
            db.close()


def _mark_preinstall_complete(db):
    """Mark preinstall as complete in SystemSetting."""
    from database import SystemSetting

    setting = db.query(SystemSetting).filter(
        SystemSetting.key == "knowledge_base_preinstalled"
    ).first()

    if setting:
        setting.value = "true"
        setting.updated_at = datetime.now()
    else:
        setting = SystemSetting(
            key="knowledge_base_preinstalled",
            value="true"
        )
        db.add(setting)

    db.commit()


def _preinstall_create_competitor(db, comp_data: CompetitorData):
    """Create a new competitor with preinstall labeling."""
    from database import Competitor, DataSource

    competitor = Competitor(
        name=comp_data.canonical_name.title(),
        website=comp_data.website,
        status=comp_data.status or "Active",
        threat_level=comp_data.threat_level or "Medium",

        # Pricing
        pricing_model=comp_data.pricing_model,
        base_price=comp_data.base_price,
        price_unit=comp_data.price_unit,

        # Product
        product_categories=comp_data.product_categories,
        key_features=comp_data.key_features,
        integration_partners=comp_data.integration_partners,
        certifications=comp_data.certifications,

        # Market
        target_segments=comp_data.target_segments,
        customer_size_focus=comp_data.customer_size_focus,
        geographic_focus=comp_data.geographic_focus,
        customer_count=comp_data.customer_count,
        customer_acquisition_rate=comp_data.customer_acquisition_rate,
        key_customers=comp_data.key_customers,

        # Company
        employee_count=comp_data.employee_count,
        employee_growth_rate=comp_data.employee_growth_rate,
        year_founded=comp_data.year_founded,
        headquarters=comp_data.headquarters,
        funding_total=comp_data.funding_total,
        latest_round=comp_data.latest_round,
        pe_vc_backers=comp_data.pe_vc_backers,

        # Digital
        website_traffic=comp_data.website_traffic,
        social_following=comp_data.social_following,
        g2_rating=comp_data.g2_rating,

        # Features
        has_pxp=comp_data.has_patient_intake,
        has_rcm=comp_data.has_insurance_verification,
        has_payments=comp_data.has_payments,
        has_patient_mgmt=comp_data.has_patient_portal,

        # Notes
        notes=comp_data.notes,
        recent_launches=comp_data.recent_launches,

        # Metadata
        created_at=datetime.now(),
        last_updated=datetime.now()
    )

    db.add(competitor)
    db.flush()  # Get the ID

    # Create DataSource records with "Preinstalled" label
    _create_preinstall_data_sources(db, competitor.id, comp_data)

    return competitor


def _preinstall_update_competitor(db, existing, comp_data: CompetitorData) -> int:
    """Update existing competitor with preinstall data (fill gaps only)."""
    from database import Competitor, DataSource

    fields_updated = 0

    field_mapping = {
        "website": "website",
        "pricing_model": "pricing_model",
        "base_price": "base_price",
        "price_unit": "price_unit",
        "product_categories": "product_categories",
        "key_features": "key_features",
        "integration_partners": "integration_partners",
        "certifications": "certifications",
        "target_segments": "target_segments",
        "customer_size_focus": "customer_size_focus",
        "geographic_focus": "geographic_focus",
        "customer_count": "customer_count",
        "customer_acquisition_rate": "customer_acquisition_rate",
        "key_customers": "key_customers",
        "employee_count": "employee_count",
        "employee_growth_rate": "employee_growth_rate",
        "year_founded": "year_founded",
        "headquarters": "headquarters",
        "funding_total": "funding_total",
        "latest_round": "latest_round",
        "pe_vc_backers": "pe_vc_backers",
        "website_traffic": "website_traffic",
        "social_following": "social_following",
        "g2_rating": "g2_rating",
        "notes": "notes",
        "recent_launches": "recent_launches",
    }

    for data_field, db_field in field_mapping.items():
        new_value = getattr(comp_data, data_field, None)
        if not new_value:
            continue

        existing_value = getattr(existing, db_field, None)

        # Fill gaps only - never overwrite existing data
        if not existing_value:
            setattr(existing, db_field, new_value)
            fields_updated += 1

            # Create DataSource for this field
            _create_single_preinstall_source(
                db, existing.id, db_field, new_value, comp_data.source_file
            )

    if fields_updated > 0:
        existing.last_updated = datetime.now()

    return fields_updated


def _create_preinstall_data_sources(db, competitor_id: int, comp_data: CompetitorData):
    """Create DataSource records for all populated fields with Preinstalled label."""

    fields_to_track = [
        ("website", comp_data.website),
        ("pricing_model", comp_data.pricing_model),
        ("base_price", comp_data.base_price),
        ("price_unit", comp_data.price_unit),
        ("product_categories", comp_data.product_categories),
        ("key_features", comp_data.key_features),
        ("integration_partners", comp_data.integration_partners),
        ("certifications", comp_data.certifications),
        ("target_segments", comp_data.target_segments),
        ("customer_size_focus", comp_data.customer_size_focus),
        ("geographic_focus", comp_data.geographic_focus),
        ("customer_count", comp_data.customer_count),
        ("customer_acquisition_rate", comp_data.customer_acquisition_rate),
        ("key_customers", comp_data.key_customers),
        ("employee_count", comp_data.employee_count),
        ("employee_growth_rate", comp_data.employee_growth_rate),
        ("year_founded", comp_data.year_founded),
        ("headquarters", comp_data.headquarters),
        ("funding_total", comp_data.funding_total),
        ("latest_round", comp_data.latest_round),
        ("pe_vc_backers", comp_data.pe_vc_backers),
        ("website_traffic", comp_data.website_traffic),
        ("social_following", comp_data.social_following),
        ("g2_rating", comp_data.g2_rating),
        ("notes", comp_data.notes),
        ("recent_launches", comp_data.recent_launches),
    ]

    for field_name, value in fields_to_track:
        if value:
            _create_single_preinstall_source(
                db, competitor_id, field_name, value, comp_data.source_file
            )


def _create_single_preinstall_source(db, competitor_id: int, field_name: str, value: str, source_file: str):
    """Create a single DataSource record with Preinstalled labeling."""
    from database import DataSource

    source = DataSource(
        competitor_id=competitor_id,
        field_name=field_name,
        current_value=value,
        source_type="client_provided",
        source_name="Certify Health (Preinstalled)",  # Clear labeling
        source_url=f"file://{source_file}",
        extraction_method="preinstall",
        confidence_score=85,  # High confidence - client provided
        confidence_level="high",
        source_reliability="B",  # Usually reliable
        information_credibility=2,  # Probably true
        is_verified=True,  # Preinstalled data is auto-verified
        verified_by="preinstall",
        verification_date=datetime.now(),
        extracted_at=datetime.now(),
        created_at=datetime.now()
    )
    db.add(source)


# ==============================================================================
# CLI TESTING
# ==============================================================================

if __name__ == "__main__":
    print("=" * 60)
    print("Certify Intel - Knowledge Base Importer")
    print("=" * 60)

    importer = KnowledgeBaseImporter()

    print(f"\nScanning folder: {importer.folder_path}")
    files = importer.scan_folder()
    print(f"Found {len(files)} supported files")

    for f in files[:10]:  # Show first 10
        print(f"  - {f.filename} ({f.file_type}, {f.size_bytes:,} bytes)")

    if len(files) > 10:
        print(f"  ... and {len(files) - 10} more")

    print("\nExtracting competitor data...")
    result = importer.extract_all()

    print(f"\n{'='*60}")
    print("IMPORT RESULTS")
    print(f"{'='*60}")
    print(f"Files scanned:      {result.total_files_scanned}")
    print(f"Files parsed:       {result.files_parsed}")
    print(f"Competitors found:  {result.competitors_found}")
    print(f"Unique competitors: {result.unique_competitors}")

    if result.errors:
        print(f"\nErrors ({len(result.errors)}):")
        for error in result.errors[:5]:
            print(f"  - {error}")

    if result.warnings:
        print(f"\nWarnings ({len(result.warnings)}):")
        for warning in result.warnings[:5]:
            print(f"  - {warning}")

    print(f"\nCompetitors extracted:")
    for comp in result.competitors[:20]:
        print(f"  - {comp.canonical_name} (from {Path(comp.source_file).name})")

    if len(result.competitors) > 20:
        print(f"  ... and {len(result.competitors) - 20} more")
