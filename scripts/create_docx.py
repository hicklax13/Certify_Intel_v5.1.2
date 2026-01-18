from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Title
title = doc.add_heading('Certify Intel - Team Task Assignments', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadata
doc.add_paragraph('Project: Certify Intel')
doc.add_paragraph('Project Start: January 5, 2026')
doc.add_paragraph('Minimal Working Prototype Due: January 22, 2026')
doc.add_paragraph('Minimal Viable Product Due: February 9, 2026')

# Part 1 Header
doc.add_heading('Part 1: Team Roles & Responsibilities', level=1)

# Product Development Lead
doc.add_heading('Product Development Lead', level=2)
doc.add_heading('Role Description', level=3)
doc.add_paragraph('The Product Development Lead is responsible for building the backend "engine" that powers the competitive intelligence system. This includes the cloud infrastructure, web scraping automation, AI-powered data extraction, and connecting the backend to the Excel dashboard via Power Query. This role ensures that raw competitor data flows automatically from websites into the Excel dashboard without manual intervention.')

doc.add_heading('Primary Focus', level=3)
doc.add_paragraph('• Building the automated data collection engine\n• Integrating GPT-4 for intelligent data extraction\n• Connecting backend to Excel dashboard\n• Ensuring weekly automated refreshes work reliably')

doc.add_heading('Minimal Working Prototype Phase Tasks (January 5-22)', level=3)
doc.add_paragraph('☐ Create Excel template structure with all 32 data columns\n☐ Define data schema and column naming conventions\n☐ Support Data Lead with technical data formatting questions\n☐ Ensure Excel file is structured for future Power Query connection')

doc.add_heading('Minimal Viable Product Phase Tasks (January 23 - February 9)', level=3)
doc.add_paragraph('☐ Set up cloud infrastructure (AWS or Azure)\n☐ Deploy PostgreSQL database for competitor data storage\n☐ Build Playwright web scraper for competitor websites\n☐ Create FastAPI backend with REST endpoints\n☐ Integrate OpenAI GPT-4 API for data extraction\n☐ Build extraction prompts for each data category\n☐ Create Power Query connection from Excel to backend API\n☐ Implement scheduled weekly scraping\n☐ Build email alert system for critical changes\n☐ Write technical documentation for backend system\n☐ Conduct backend testing and debugging')

# Team Lead
doc.add_heading('Team Lead / Coordinator', level=2)
doc.add_heading('Role Description', level=3)
doc.add_paragraph('The Team Lead serves as the project\'s primary coordinator and the main point of contact between the client, advisors, and team members. This role ensures workstreams are aligned, removes blockers, facilitates decision-making, and keeps the project on track.')

doc.add_heading('Primary Focus', level=3)
doc.add_paragraph('• Cross-team coordination and communication\n• Client and advisor relationship management\n• Decision-making and approvals\n• Ensuring deliverables meet quality standards')

doc.add_heading('Minimal Working Prototype Phase Tasks (January 5-22)', level=3)
doc.add_paragraph('☐ Kick off project with team alignment meeting\n☐ Approve "middle market" definition and competitor criteria\n☐ Review and prioritize initial competitor list\n☐ Facilitate communication with client stakeholders\n☐ Review and approve Minimal Working Prototype dashboard')

doc.add_heading('Minimal Viable Product Phase Tasks (January 23 - February 9)', level=3)
doc.add_paragraph('☐ Conduct weekly status check-in meetings\n☐ Coordinate between team members when dependencies arise\n☐ Communicate progress updates to client and advisors\n☐ Review AI-generated insights for quality\n☐ Attend Executive Presentation Workshop (January 27)\n☐ Lead intensive weekend sessions (February 6-8)\n☐ Conduct final review and sign-off\n☐ Lead final presentation to client')

# Research Lead
doc.add_heading('Research Lead', level=2)
doc.add_heading('Role Description', level=3)
doc.add_paragraph('The Research Lead is responsible for all competitive research activities. This role maps the competitive landscape, identifies who qualifies as a competitor, understands what solutions target customers currently use instead of Certify Health, and finds reliable data sources for each data point.')

doc.add_heading('Primary Focus', level=3)
doc.add_paragraph('• Defining the competitive landscape\n• Identifying and qualifying competitors\n• Understanding target customer buying behavior\n• Finding and validating data sources\n• Matching Certify Health products to competitor solutions')

doc.add_heading('Minimal Working Prototype Phase Tasks (January 5-22)', level=3)
doc.add_paragraph('☐ Define criteria for what qualifies as a competitor\n☐ Create competitor qualification framework\n☐ Identify initial list of 40+ competitors\n☐ Map competitors to Certify Health product categories\n☐ Gather website URLs and key pages for each competitor\n☐ Research target customer segments\n☐ Initial data collection for 15 priority competitors')

doc.add_heading('Minimal Viable Product Phase Tasks (January 23 - February 9)', level=3)
doc.add_paragraph('☐ Complete competitive landscape map\n☐ Identify data sources for each of the 32 data points\n☐ Document where to find pricing information\n☐ Document where to find customer/logo information\n☐ Research funding and valuation data\n☐ Research hiring trends and employee counts\n☐ Validate AI-extracted data for accuracy\n☐ Monitor competitor news weekly\n☐ Create research methodology documentation')

# Data Lead
doc.add_heading('Data Lead', level=2)
doc.add_heading('Role Description', level=3)
doc.add_paragraph('The Data Lead is responsible for building and maintaining the Excel dashboard. This role takes the research from the Research Lead and translates it into structured data points that populate the dashboard.')

doc.add_heading('Primary Focus', level=3)
doc.add_paragraph('• Building the Excel dashboard structure\n• Translating research into standardized data points\n• Creating KPI formulas and calculations\n• Ensuring data accuracy and consistency')

doc.add_heading('Minimal Working Prototype Phase Tasks (January 5-22)', level=3)
doc.add_paragraph('☐ Build Excel workbook structure with all sheets\n☐ Create data entry sheet with 32 columns\n☐ Define data types and validation rules\n☐ Populate 10-15 competitors with research data\n☐ Build initial pivot tables for data analysis\n☐ Create dropdown lists for standardized fields')

doc.add_heading('Minimal Viable Product Phase Tasks (January 23 - February 9)', level=3)
doc.add_paragraph('☐ Expand competitor data to all 40+ companies\n☐ Build pricing comparison matrix\n☐ Build feature comparison matrix\n☐ Create KPI formulas\n☐ Build market segment analysis views\n☐ Create conditional formatting rules\n☐ Test Power Query data refresh\n☐ Conduct data quality audit\n☐ Document data definitions and sources')

# Delivery Lead
doc.add_heading('Delivery Lead', level=2)
doc.add_heading('Role Description', level=3)
doc.add_paragraph('The Delivery Lead is responsible for the visual presentation and documentation of the dashboard. This role focuses on making the dashboard visually compelling, easy to understand, and presentation-ready.')

doc.add_heading('Primary Focus', level=3)
doc.add_paragraph('• Dashboard layout and visual design\n• Charts, graphs, and data visualizations\n• User experience and navigation\n• Documentation and user guides')

doc.add_heading('Minimal Working Prototype Phase Tasks (January 5-22)', level=3)
doc.add_paragraph('☐ Design overall dashboard layout and navigation\n☐ Create color scheme and visual style guide\n☐ Design chart templates\n☐ Build dashboard summary/overview sheet\n☐ Apply professional formatting and branding')

doc.add_heading('Minimal Viable Product Phase Tasks (January 23 - February 9)', level=3)
doc.add_paragraph('☐ Create advanced visualizations\n☐ Design competitor profile view layout\n☐ Build threat level visualization\n☐ Create market segment breakdown charts\n☐ Write user guide documentation\n☐ Create "How to Use" instructions sheet\n☐ Polish final dashboard visuals\n☐ Prepare presentation materials\n☐ Support Clinics Intensive Weekend')

# Part 2
doc.add_heading('Part 2: Key Dates Summary', level=1)
doc.add_paragraph('January 5, 2026 - Project Start / Team Onboarding\nJanuary 22, 2026 - Minimal Working Prototype Due\nJanuary 27, 2026 - Executive Presentation Workshop\nFebruary 6-8, 2026 - Clinics Intensive Weekend\nWeek of February 9, 2026 - Final Presentations / Minimal Viable Product Due')

# Save
doc.save(r'C:\Users\conno\Downloads\Certify_Health_Intelv1\Certify_Health_Team_1_Task_Assignments.docx')
print('Word document created successfully!')
